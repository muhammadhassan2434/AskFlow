from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.exceptions.document import (
    CorruptedDocxError,
    EmptyParsedDocumentError,
)
from app.models.download_result import DownloadResult
from app.models.parser_result import ParserResult
from app.parsers.base_parser import BaseParser


class DOCXParser(BaseParser):
    """
    Parses Microsoft Word (.doc/.docx) documents.

    Version 1 supports:
    - Paragraphs
    - Tables

    Future versions may support:
    - Images (OCR)
    - Headers & Footers
    - Comments
    - Track Changes
    """

    def parse(
        self,
        download: DownloadResult,
    ) -> ParserResult:

        try:

            document = Document(download.local_path)

        except PackageNotFoundError as exc:

            raise CorruptedDocxError(
                "Invalid or corrupted DOCX document."
            ) from exc

        except Exception as exc:

            raise CorruptedDocxError(
                "Unable to open DOCX document."
            ) from exc

        sections: list[str] = []

        # ----------------------------------
        # Extract paragraphs
        # ----------------------------------

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:

                sections.append(text)

        # ----------------------------------
        # Extract tables
        # ----------------------------------

        for table in document.tables:

            for row in table.rows:

                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if cells:

                    sections.append(" | ".join(cells))

        full_text = "\n\n".join(sections).strip()

        if not full_text:

            raise EmptyParsedDocumentError(
                "No readable text found in DOCX document."
            )

        return ParserResult(
            text=full_text,
            metadata={
                "paragraphs": len(document.paragraphs),
                "tables": len(document.tables),
            },
        )